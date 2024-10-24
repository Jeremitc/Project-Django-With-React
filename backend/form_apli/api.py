from rest_framework import viewsets, permissions
from .models import ReciboProvisional
from .serializers import ReciboProvisionalSerializer
from rest_framework.decorators import action
from rest_framework.response import Response
from rest_framework import status
from docx import Document
from django.http import HttpResponse

class form_apliViewSet(viewsets.ModelViewSet):
    queryset = ReciboProvisional.objects.all()
    permission_classes = [permissions.AllowAny]
    serializer_class = ReciboProvisionalSerializer

    @action(detail=True, methods=['POST'])
    def done(self, request, pk=None):
        recibo_provisional = self.get_object()
        recibo_provisional.save()
        
        # Generate the Word document
        doc = Document()
        doc.add_heading('CONTRATO DE ARRENDAMIENTO DE LOCAL COMERCIAL', 0)

        # Replace the text with data from the model
        doc.add_paragraph(f"Arrendador: {recibo_provisional.nombre}")
        doc.add_paragraph(f"DNI: {recibo_provisional.dni}")
        doc.add_paragraph(f"RUC: {recibo_provisional.ruc}")
        doc.add_paragraph(f"Dirección: {recibo_provisional.direccion}")
        doc.add_paragraph(f"Local: {recibo_provisional.local}")
        doc.add_paragraph(f"Área: {recibo_provisional.area} m2")
        doc.add_paragraph(f"Uso: {recibo_provisional.uso_exclusivo}")
        doc.add_paragraph(f"Renta mensual: ${recibo_provisional.renta_mensual}")
        doc.add_paragraph(f"Garantía: ${recibo_provisional.garantia}")
        doc.add_paragraph(f"Fecha de inicio: {recibo_provisional.fecha_inicio}")

        # Send the document as a response
        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename=contrato_{recibo_provisional.id}.docx'
        doc.save(response)

        return response
